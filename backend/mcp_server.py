# mcp_server.py
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os

WORKSPACE_DOC = os.path.join(os.path.dirname(__file__), "workspace", "current.docx")
TOOL_MAP = {}

def _get_doc():
    if not os.path.exists(WORKSPACE_DOC):
        raise FileNotFoundError("工作区文档不存在")
    return Document(WORKSPACE_DOC)

def _save_doc(doc):
    doc.save(WORKSPACE_DOC)

def _get_paragraph(para_index, doc):
    if para_index < 0 or para_index >= len(doc.paragraphs):
        raise IndexError(f"段落索引 {para_index} 无效，文档共有 {len(doc.paragraphs)} 段")
    return doc.paragraphs[para_index]

# ---------- 工具函数（不再用 @server.tool()）----------
async def set_font(para_index: int, font_east_asia: str = None,
                   font_ascii: str = None, font_size: str = None):
    """修改整个段落的所有 Run 字体和字号"""
    doc = _get_doc()
    para = _get_paragraph(para_index, doc)
    for run in para.runs:
        if font_east_asia:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = rPr.makeelement(qn('w:rFonts'), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), font_east_asia)
        if font_ascii:
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = rPr.makeelement(qn('w:rFonts'), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), font_ascii)
        if font_size is not None:
            # 兼容数字和字符串
            if isinstance(font_size, str):
                pt_val = float(font_size.replace('pt', ''))
            else:
                pt_val = float(font_size)
            run.font.size = Pt(pt_val)
    _save_doc(doc)
    return f"段落 {para_index} 字体已修改"

TOOL_MAP["set_font"] = set_font

async def set_bold(para_index: int, bold: bool = True):
    doc = _get_doc()
    para = _get_paragraph(para_index, doc)
    for run in para.runs:
        run.bold = bold
    _save_doc(doc)
    return f"段落 {para_index} 加粗: {bold}"

TOOL_MAP["set_bold"] = set_bold

async def set_alignment(para_index: int, alignment: str = "center"):
    doc = _get_doc()
    para = _get_paragraph(para_index, doc)
    align_map = {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    }
    if alignment not in align_map:
        return f"不支持的对齐方式: {alignment}"
    para.alignment = align_map[alignment]
    _save_doc(doc)
    return f"段落 {para_index} 对齐方式: {alignment}"

TOOL_MAP["set_alignment"] = set_alignment

async def set_indent(para_index: int, first_line_indent: str = None,
                     left_indent: str = None, right_indent: str = None):
    doc = _get_doc()
    para = _get_paragraph(para_index, doc)
    pf = para.paragraph_format
    if first_line_indent:
        if first_line_indent.endswith("ch"):
            ch = float(first_line_indent.replace("ch", ""))
            pf.first_line_indent = Cm(ch * 0.37)
        else:
            pf.first_line_indent = Cm(float(first_line_indent.replace("cm", "")))
    if left_indent:
        pf.left_indent = Cm(float(left_indent.replace("cm", "")))
    if right_indent:
        pf.right_indent = Cm(float(right_indent.replace("cm", "")))
    _save_doc(doc)
    return f"段落 {para_index} 缩进已修改"

TOOL_MAP["set_indent"] = set_indent

#行距和段间距设置，支持倍数和固定值两种模式
async def set_spacing(para_index: int, line_spacing: str = None,
                      space_before: str = None, space_after: str = None):
    """
    设置段间距和行距。
    line_spacing: 数字(如1.5)表示倍数；字符串"20pt"表示固定值。
    """
    doc = _get_doc()
    para = _get_paragraph(para_index, doc)
    pf = para.paragraph_format

    if line_spacing is not None:
        if isinstance(line_spacing, str) and line_spacing.endswith('pt'):
            # 固定行距
            pt_val = float(line_spacing.replace('pt', ''))
            pf.line_spacing = Pt(pt_val)
        else:
            # 倍数
            pf.line_spacing = float(line_spacing)

    if space_before:
        pf.space_before = Pt(float(space_before.replace('pt', '')))
    if space_after:
        pf.space_after = Pt(float(space_after.replace('pt', '')))

    _save_doc(doc)
    return f"段落 {para_index} 间距已设置"

TOOL_MAP["set_spacing"] = set_spacing

async def set_color(para_index: int, color: str):
    doc = _get_doc()
    para = _get_paragraph(para_index, doc)
    if color.startswith("#"):
        r, g, b = bytes.fromhex(color[1:])
        rgb = RGBColor(r, g, b)
    else:
        color_map = {
            "red": RGBColor(255, 0, 0),
            "blue": RGBColor(0, 0, 255),
            "green": RGBColor(0, 255, 0),
            "black": RGBColor(0, 0, 0)
        }
        rgb = color_map.get(color.lower(), RGBColor(0, 0, 0))
    for run in para.runs:
        run.font.color.rgb = rgb
    _save_doc(doc)
    return f"段落 {para_index} 颜色已修改"

TOOL_MAP["set_color"] = set_color