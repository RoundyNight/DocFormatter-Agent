# doc_tools.py -- LangChain @tool 定义，用于 bind_tools + ToolNode
import os
import tempfile
import threading
from typing import Optional, List
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from langchain_core.tools import tool

WORKSPACE_DOC = os.path.join(os.path.dirname(__file__), "workspace", "current.docx")
_doc_lock = threading.Lock()

# ---------- 内部辅助函数 ----------
def _save_doc(doc: Document) -> None:
    """原子写入：先写临时文件再 replace，避免并发/同路径保存损坏 zip。"""
    dest_dir = os.path.dirname(WORKSPACE_DOC)
    fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=dest_dir)
    os.close(fd)
    try:
        doc.save(tmp_path)
        os.replace(tmp_path, WORKSPACE_DOC)
    except Exception:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

def _get_paragraph(idx: int, doc: Document):
    if idx < 0 or idx >= len(doc.paragraphs):
        raise IndexError(f"段落索引 {idx} 无效，文档共有 {len(doc.paragraphs)} 段")
    return doc.paragraphs[idx]

def _parse_indices(para_index) -> list:
    """兼容 int, list, 或 '1,2,3' 字符串格式的段落索引"""
    if isinstance(para_index, int):
        return [para_index]
    if isinstance(para_index, list):
        return [int(i) for i in para_index]
    if isinstance(para_index, str):
        return [int(i.strip()) for i in para_index.replace(" ", "").split(",") if i.strip()]
    raise ValueError(f"无效的 para_index 格式: {para_index}")

def _ensure_runs(para):
    """无 Run 的段落无法设置字体/加粗，补一个 Run。"""
    if not para.runs:
        para.add_run("")

# ---------- LangChain @tool 定义 ----------

@tool
def set_font(
    para_index: list[int],
    font_east_asia: Optional[str] = None,
    font_ascii: Optional[str] = None,
    font_size: Optional[str] = None,
) -> str:
    """批量修改段落的字体和字号。para_index 为段落索引数组，如 [0,1,2]。
    font_east_asia: 中文字体名，如"黑体"、"宋体"；font_ascii: 西文字体名；font_size: 字号字符串，如"14"或"14pt"。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            _ensure_runs(para)
            for run in para.runs:
                if font_east_asia:
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn("w:rFonts"), {})
                        rPr.insert(0, rFonts)
                    rFonts.set(qn("w:eastAsia"), font_east_asia)
                if font_ascii:
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn("w:rFonts"), {})
                        rPr.insert(0, rFonts)
                    rFonts.set(qn("w:ascii"), font_ascii)
                if font_size is not None:
                    pt_val = float(str(font_size).replace("pt", ""))
                    run.font.size = Pt(pt_val)
        _save_doc(doc)
    return f"段落 {indices} 字体已修改"

@tool
def set_bold(
    para_index: list[int],
    bold: bool = True,
) -> str:
    """批量设置段落加粗。para_index 为段落索引数组；bold 为 True 加粗，False 取消加粗。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            _ensure_runs(para)
            for run in para.runs:
                run.bold = bold
        _save_doc(doc)
    return f"段落 {indices} 加粗: {bold}"

@tool
def set_alignment(
    para_index: list[int],
    alignment: str = "center",
) -> str:
    """批量设置段落对齐方式。para_index 为段落索引数组；alignment 取值: left, center, right, justify。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        align_map = {
            "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        if alignment not in align_map:
            return f"不支持的对齐方式: {alignment}"
        for idx in indices:
            para = _get_paragraph(idx, doc)
            para.alignment = align_map[alignment]
        _save_doc(doc)
    return f"段落 {indices} 对齐方式: {alignment}"

@tool
def set_indent(
    para_index: list[int],
    first_line_indent: Optional[str] = None,
    left_indent: Optional[str] = None,
    right_indent: Optional[str] = None,
) -> str:
    """批量设置段落缩进。para_index 为段落索引数组。
    first_line_indent: 首行缩进，如"2ch"(2字符)或"0.74cm"；left_indent/right_indent: 左/右缩进，单位cm。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            pf = para.paragraph_format
            if first_line_indent:
                if first_line_indent.endswith("ch"):
                    pf.first_line_indent = Cm(float(first_line_indent.replace("ch", "")) * 0.37)
                else:
                    pf.first_line_indent = Cm(float(first_line_indent.replace("cm", "")))
            if left_indent:
                pf.left_indent = Cm(float(left_indent.replace("cm", "")))
            if right_indent:
                pf.right_indent = Cm(float(right_indent.replace("cm", "")))
        _save_doc(doc)
    return f"段落 {indices} 缩进已修改"

@tool
def set_spacing(
    para_index: list[int],
    line_spacing: Optional[str] = None,
    space_before: Optional[str] = None,
    space_after: Optional[str] = None,
) -> str:
    """批量设置段落间距。para_index 为段落索引数组。
    line_spacing: 行距(数字=倍数如1.5, 字符如"20pt"=固定值)；space_before/space_after: 段前/段后距，如"10pt"。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        for idx in indices:
            para = _get_paragraph(idx, doc)
            pf = para.paragraph_format
            if line_spacing is not None:
                if isinstance(line_spacing, str) and line_spacing.endswith("pt"):
                    pf.line_spacing = Pt(float(line_spacing.replace("pt", "")))
                else:
                    pf.line_spacing = float(line_spacing)
            if space_before:
                pf.space_before = Pt(float(space_before.replace("pt", "")))
            if space_after:
                pf.space_after = Pt(float(space_after.replace("pt", "")))
        _save_doc(doc)
    return f"段落 {indices} 间距已设置"

@tool
def set_color(
    para_index: list[int],
    color: str,
) -> str:
    """批量设置段落字体颜色。para_index 为段落索引数组；color 取值: #hex 如 #FF0000，或名称如 red/blue/green/black。"""
    with _doc_lock:
        doc = Document(WORKSPACE_DOC)
        indices = _parse_indices(para_index)
        if color.startswith("#"):
            r, g, b = bytes.fromhex(color[1:])
            rgb = RGBColor(r, g, b)
        else:
            color_map = {
                "red": RGBColor(255, 0, 0),
                "blue": RGBColor(0, 0, 255),
                "green": RGBColor(0, 255, 0),
                "black": RGBColor(0, 0, 0),
            }
            rgb = color_map.get(color.lower(), RGBColor(0, 0, 0))
        for idx in indices:
            para = _get_paragraph(idx, doc)
            _ensure_runs(para)
            for run in para.runs:
                run.font.color.rgb = rgb
        _save_doc(doc)
    return f"段落 {indices} 颜色已修改"

# ---------- 工具列表 (供 bind_tools 和 execute_tools 使用) ----------
all_tools = [set_font, set_bold, set_alignment, set_indent, set_spacing, set_color]
